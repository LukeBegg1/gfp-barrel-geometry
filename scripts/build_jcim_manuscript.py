#!/usr/bin/env python3
"""
Build JCIM-style manuscript: Introduction, Methods, Results, Discussion.
ACS numbered references. Figures inserted inline.
"""
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
fig_dir = '/Users/lukebegg/Downloads/deep_analysis/pub_figures'

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 2.0

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run(
    'AI-Assisted Systematic Analysis of Beta-Barrel Cross-Sectional '
    'Geometry Across 908 Fluorescent Protein Crystal Structures: '
    'Relationships Between Barrel Shape, Chromophore Rigidity, and '
    'Photophysical Properties'
)
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'
doc.add_paragraph()

authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = authors.add_run('[Author Names]')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

affil = doc.add_paragraph()
affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = affil.add_run('[Affiliations]')
run.font.size = Pt(11)
run.font.name = 'Times New Roman'
run.italic = True
doc.add_paragraph()

# Abstract
add_abs_head = doc.add_paragraph()
add_abs_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = add_abs_head.add_run('Abstract')
r.bold = True
r.font.size = Pt(12)
r.font.name = 'Times New Roman'

abs_p = doc.add_paragraph()
abs_p.paragraph_format.line_spacing = 2.0
abs_text = (
    'The 11-stranded beta-barrel of fluorescent proteins is universally '
    'conserved, yet its quantitative geometry has not been systematically '
    'characterized across the structural database. Here, we report a '
    'computational analysis of beta-barrel cross-sectional geometry across '
    '908 fluorescent protein crystal structures from the RCSB Protein Data '
    'Bank. For each structure, the barrel axis was determined by principal '
    'component analysis of backbone C\u03b1 coordinates, and a '
    'perpendicular cross-sectional slice at the chromophore level was '
    'characterized by convex hull analysis. The entire computational '
    'pipeline\u2014including CIF parsing, coordinate transformation, '
    'geometry quantification, spectral data integration, and statistical '
    'analysis\u2014was developed using Claude (Anthropic), a large language '
    'model accessed through the agentic coding environment Claude Code. '
    'Barrel shape, but not size, correlates with emission wavelength: '
    'red-shifted proteins have narrower, more elliptical barrels '
    '(\u03c1 = \u20130.290 for minor axis, p < 10\u207b\u00b9\u2074). '
    'The ratio of chromophore to barrel crystallographic B-factors is the '
    'strongest predictor of quantum yield (\u03c1 = \u20130.493, '
    'p < 10\u207b\u00b9\u2077). Chromophore dihedral angles defined by '
    'Megley et al. correlate with both barrel eccentricity and '
    'photophysical performance. The principal correlations survive '
    'Benjamini\u2013Hochberg correction for multiple testing and partial '
    'correlation analysis controlling for crystallographic resolution. '
    'These findings suggest that the barrel is not a passive scaffold but '
    'participates in modulating chromophore rigidity and photophysical '
    'output.'
)
r = abs_p.add_run(abs_text)
r.font.size = Pt(12)
r.font.name = 'Times New Roman'

doc.add_paragraph()

# Graphical Abstract (TOC Graphic)
toc_label = doc.add_paragraph()
toc_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = toc_label.add_run('TOC Graphic / Graphical Abstract')
r.bold = True
r.font.size = Pt(12)
r.font.name = 'Times New Roman'

ga_path = os.path.join(fig_dir, 'graphical_abstract.png')
if os.path.exists(ga_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(ga_path, width=Inches(5.5))

doc.add_paragraph()
doc.add_paragraph()


def add_heading(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

def add_subheading(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

def add_body(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    return p

def add_body_mixed(parts):
    """Add a paragraph with mixed formatting. parts is a list of (text, bold, italic) tuples."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.line_spacing = 2.0
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run.bold = bold
        run.italic = italic
    return p

def add_figure(filename, caption, width=6.5):
    path = os.path.join(fig_dir, filename)
    if os.path.exists(path):
        doc.add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(path, width=Inches(width))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cap.paragraph_format.space_after = Pt(6)
        run = cap.add_run(caption)
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        doc.add_paragraph()


# ═══════════════════════════════════════════════════════════
# INTRODUCTION
# ═══════════════════════════════════════════════════════════
add_heading('Introduction')

add_body_mixed([
    ('Green fluorescent protein (GFP), first observed in the jellyfish ', False, False),
    ('Aequorea victoria', False, True),
    (' by Shimomura et al. in 1962,\u00b9 has fundamentally '
     'transformed biological research by enabling the direct visualization of '
     'proteins, organelles, and cellular processes in living systems.\u00b2 The '
     'subsequent cloning and heterologous expression of GFP by Chalfie et al.\u00b3 '
     'demonstrated that the protein fluoresces without requiring cofactors or '
     'external enzymes, establishing GFP as a genetically encodable marker. '
     'The significance of this discovery was recognized with the 2008 Nobel '
     'Prize in Chemistry, awarded to Shimomura, Chalfie, and Tsien.\u2074 In the '
     'decades since, fluorescent proteins have become indispensable tools '
     'across cell biology, neuroscience, and biotechnology, serving as '
     'reporters, biosensors, and partners in FRET\u2075 and, more recently, '
     'super-resolution microscopy.\u2076\u02c9\u2077', False, False),
])

add_body(
    'The three-dimensional structure of GFP, resolved by X-ray '
    'crystallography in 1996, revealed a distinctive 11-stranded beta-barrel '
    'fold approximately 42 \u00c5 in length and 24 \u00c5 in diameter.\u2078\u02c9\u2079 '
    'This barrel encases a central alpha-helix, with the chromophore\u2014formed '
    'by autocatalytic post-translational cyclization of residues '
    'Ser65\u2013Tyr66\u2013Gly67\u2014positioned at the geometric center, '
    'shielded from solvent.\u00b9\u2070\u02c9\u00b9\u00b9 The rigid encapsulation '
    'minimizes non-radiative decay, enabling high quantum yields. The same '
    '11-stranded topology is conserved across all known fluorescent protein '
    'families, including GFP-derived variants and fluorescent proteins from '
    'anthozoan organisms such as DsRed, despite sequence identities as low as 25%.\u00b9\u00b2\u02c9\u00b9\u00b3'
)

add_body_mixed([
    ('Protein engineering has produced a palette of fluorescent protein '
     'variants spanning the visible spectrum. Substitution of Tyr66 with '
     'tryptophan or histidine yields cyan or blue variants, '
     'respectively.\u00b9\u2074 The Thr203Tyr mutation introduces a \u03c0-stacking '
     'interaction that red-shifts emission by ~20 nm.\u00b9\u2075 In the DsRed '
     'family, an additional oxidation extends the conjugated \u03c0-system, '
     'producing red emission.\u00b9\u2076 Structure-guided engineering has '
     'achieved notable results: Goedhart et al. obtained a '
     'quantum yield of 0.93 in mTurquoise2 through a single packing mutation '
     '(I146F) identified by structure-guided saturation mutagenesis,\u00b9\u2077 while Campbell and '
     'colleagues have systematically monomerized and optimized coral-derived '
     'variants.\u00b9\u2078\u02c9\u00b9\u2079 Hirano et al. developed StayGold, '
     'a green fluorescent protein from the jellyfish ', False, False),
    ('Cytaeis uchidae', False, True),
    (', which exhibits exceptional photostability.\u2074\u00b9 A persistent trend '
     'across this engineering is that quantum yield decreases with increasing '
     'emission wavelength.\u00b2\u2070 This has been attributed to increased '
     'vibrational degrees of freedom in extended conjugation systems, which '
     'open additional non-radiative relaxation pathways.', False, False),
])

add_body(
    'Despite the wealth of structural data now available\u2014over 800 crystal '
    'structures classified within the GFP-like superfamily by SCOP, CATH, and '
    'Pfam\u2014the quantitative geometry of the beta-barrel itself has received '
    'little systematic attention. Previous studies have focused on chromophore '
    'chemistry, hydrogen bonding networks, and site-specific mutation effects '
    'on spectral properties.\u00b2\u00b9\u207b\u00b2\u00b3 Computational work by '
    'Zimmer and colleagues has shown that the barrel contracts upon '
    'chromophore maturation and that water permeability through the barrel '
    'wall varies between GFP variants.\u00b2\u2074\u02c9\u00b2\u2075 Analysis of '
    'conserved glycine residues at positions 31, 33, and 35 demonstrated '
    'that these residues are essential for proper barrel folding and '
    'assembly.\u00b2\u2076 Megley et al. '
    'showed that the \u03c6 and \u03c4 dihedral angles of the chromophore '
    'differ between yellow, blue, and green variants and correlate with '
    'non-radiative decay.\u00b2\u2077 Ahmed et al. used molecular dynamics to '
    'guide engineering of YuzuFP, identifying a mutation at residue 148 that '
    'increased brightness 1.5-fold.\u00b2\u2078 However, the beta-barrel has '
    'generally been treated as a static scaffold. Whether barrel geometry '
    'varies systematically across fluorescent protein families, and whether '
    'such variation relates to photophysical properties, has not been '
    'addressed.'
)

add_body(
    'Here, we report a systematic analysis of beta-barrel cross-sectional '
    'geometry across 908 fluorescent protein crystal structures from the '
    'RCSB Protein Data Bank. For each structure, the barrel axis was '
    'determined by PCA of backbone C\u03b1 coordinates, the structure was '
    'rotated to align this axis with the z-axis, and a perpendicular '
    'cross-sectional slice was extracted at the chromophore level. Convex '
    'hull analysis yielded area, eccentricity, circularity, and axis lengths '
    'for each barrel. Chromophore \u03c4 and \u03c6 dihedral angles as '
    'defined by Megley et al.\u00b2\u2077 were extracted and incorporated '
    'into the analysis. These measurements were integrated with spectral '
    'data from FPbase,\u00b3\u2074 crystallographic B-factors, and '
    'chromophore\u2013barrel contact counts. The results show that emission '
    'wavelength varies with barrel shape, that chromophore rigidity '
    'relative to the barrel is a strong predictor of quantum yield '
    '(\u03c1 = \u20130.493), and that dihedral twist correlates with '
    'both barrel eccentricity and photophysical performance.'
)

add_body(
    'This study was conducted with the assistance of Claude (Anthropic), '
    'a large language model accessed through Claude Code, an agentic coding '
    'environment that provides the model with file system access, terminal '
    'commands, and persistent memory across tasks.\u00b2\u2079 The project was '
    'organized as a sequence of discrete computational tasks: dataset '
    'assembly, chromophore detection, PCA-based barrel axis determination, '
    'convex hull analysis, spectral data matching, B-factor extraction, '
    'contact counting, torsion angle computation, and statistical analysis. '
    'For each task, the authors specified the objective and the model '
    'generated Python code, which was then reviewed, tested against known '
    'structures, and validated before application to the full dataset. This '
    'workflow shares features with recent work by Schwartz,\u00b3\u2070 who '
    'used Claude Code for theoretical physics calculations by writing '
    'intermediate results to disk for retrieval across tasks. A similar '
    'strategy was adopted here: each computational step wrote its output to '
    'CSV files on disk, which were read back in subsequent steps, allowing '
    'the full dataset to be processed across multiple sessions '
    'without requiring the model to hold all results in context.'
)

add_body(
    'The use of AI in fluorescent protein research has taken several forms. '
    'Hartley et al. used AlphaFold2 and RoseTTAFold to predict '
    'post-translational modifications in GFP-like proteins.\u00b3\u00b9 '
    'The ESM3 protein language model was used to design esmGFP, a functional '
    'fluorescent protein with only 58% identity to known GFPs;\u00b3\u00b2 '
    'related protein-specific transformers have predicted fluorescence '
    'properties from sequence for directed FP engineering.\u00b3\u00b3 '
    'These approaches use models trained on protein sequence or structural '
    'data for design tasks. By contrast, this study employs a general-purpose '
    'conversational LLM (Claude, Anthropic) as a research assistant for '
    'systematic crystallographic and statistical analysis of existing '
    'structural data\u2014a distinct role not represented in the prior '
    'FP literature. To our knowledge, this is among the first applications '
    'of a conversational LLM in this analytical role within fluorescent '
    'protein structural research. '
    'We describe the specific tasks performed with AI assistance, the '
    'failure modes encountered, and a detailed case study of systematic '
    'chromophore mis-detection and its correction in the Methods section '
    'and Supporting Information (Note S1). All scientific decisions, '
    'interpretation, and responsibility for the claims in this paper rest '
    'with the human authors.'
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════
# METHODS
# ═══════════════════════════════════════════════════════════
add_heading('Methods')

add_subheading('Dataset Assembly.')
add_body(
    'The RCSB Protein Data Bank was queried for all X-ray crystal structures '
    'belonging to the GFP-like superfamily using annotations from SCOP '
    '(b.67.1), CATH (2.60.120.200), Pfam (PF01353), and InterPro '
    '(IPR011584, IPR000786). The union returned 885 entries. Coordinate '
    'files in mmCIF format were downloaded for all entries. Eight structures '
    'were removed after inspection (APOBEC3H complexes, a DNA co-crystal, '
    'and misannotated entries). Sequences were extracted and aligned using '
    'Clustal Omega\u00b3\u2075 and inspected in Jalview\u00b3\u2076 to confirm '
    'GFP/DsRed motifs. The initial dataset comprised 877 structures; '
    'cross-referencing against additional RCSB annotations identified '
    '31 confirmed FP barrels deposited after the last SCOP update, '
    'yielding a final dataset of 908 structures (Table S1).'
)

add_subheading('Chromophore Detection.')
add_body(
    'Each structure was classified by scanning non-standard residues for '
    'known chromophore residue names. The detection list encompasses all '
    'major fluorescent chromophore variants, requiring two criteria: '
    '(i) an imidazolinone ring (atoms N2, C2, O2 in the PDB CONECT '
    'record), and (ii) a pi-conjugated aromatic ring at position 66 '
    '(phenol from Tyr66, indole from Trp66, or imidazole from His66). '
    'Criterion (ii) is biochemically necessary for fluorescence: without '
    'an aromatic donor at position 66, the push-pull conjugation '
    'responsible for visible absorption is absent. Detection covered '
    'Tyr66-based GFP-type chromophores (CRO, CR2, CRQ, GYC, CRG, NYG, '
    'and others), Trp66-based CFP-type chromophores (PIA, CRF, CRK, '
    'CR7, CFY, CCY, and others), His66-based BFP-type chromophores '
    '(GYS, IIC, CSH, and others), and red/orange-shifted variants '
    '(NRQ, CH6, CH7, SWG, OHD, RC7, B2H, C12, XYG, DYG, and others; '
    '72 distinct residue codes in total, including CR8, 4F3, and VYA '
    'whose aromatic rings use non-standard PDB atom names). '
    'Residues with an intact imidazolinone ring but lacking any '
    'aromatic ring system at position 66 (CRX, MDO, C99, CLV, '
    'KWS, Q2K, NRP, and related synthetic analogs without a '
    'phenol, indole, or imidazole donor) were excluded as '
    'non-fluorescent. Structures in which the chromophore precursor '
    'is stored as standard residues SER65\u2013TYR66\u2013GLY67 '
    '(uncyclized/immature state) were likewise classified as lacking '
    'a mature chromophore. '
    'Of 908 structures, 843 contained a validated mature chromophore '
    'and 65 did not (Table S1).'
)

add_subheading('Chromophore Torsion Angles.')
add_body(
    'Chromophore dihedral angles were defined following Megley et al.:\u00b2\u2077 '
    '\u03c4 = N\u2081\u2013C\u2081\u2013C\u2082\u2013C\u2083 (PDB atoms '
    'N2\u2013CA2\u2013CB2\u2013CG2), describing rotation about the '
    'imidazolinone\u2013bridge bond, and '
    '\u03c6 = C\u2081\u2013C\u2082\u2013C\u2083\u2013C\u2084 (atoms '
    'CA2\u2013CB2\u2013CG2\u2013CD1), describing rotation about the '
    'bridge\u2013phenol bond. Structures with |\u03c4| < 30\u00b0 were '
    'classified as cis, |\u03c4| > 150\u00b0 as trans, and the remainder as '
    'twisted. Valid angles were obtained for 690 structures (Table S1).'
)

add_subheading('Barrel Axis Determination.')
add_body(
    'The barrel axis was determined independently for each structure by PCA '
    'of backbone C\u03b1 coordinates. The eigenvector corresponding to the '
    'largest eigenvalue (PC1) was taken as the barrel axis. The eigenvalue '
    'ratio \u03bb\u2081/\u03bb\u2082 averaged 2.18, confirming elongated '
    'barrel geometry. A rotation matrix mapping PC1 onto [0, 0, 1] was '
    'applied to all non-hydrogen coordinates. The cross-sectional slice was '
    'centered on the chromophore (mean z of chromophore atoms = 0) and '
    'defined as all non-hydrogen atoms with |z| \u2264 2.0 \u00c5.'
)

add_subheading('Geometry Quantification.')
add_body(
    'All non-hydrogen atomic positions in each slice were projected onto the '
    'xy-plane (i.e., the plane perpendicular to the barrel axis). The '
    'convex hull was computed using scipy.spatial.ConvexHull.\u00b3\u2077 '
    'Cross-sectional area, perimeter, and an ellipse fit (via the 2D '
    'covariance matrix) yielded major and minor axis lengths. Eccentricity '
    'was calculated as e = \u221a(1 \u2013 b\u00b2/a\u00b2). Circularity '
    'was defined as 4\u03c0A/P\u00b2. A ring-shape test confirmed that '
    'slices captured the barrel wall (peripheral:central density ratios of '
    '5.4\u20137.0). '
    'These measurements describe the outer envelope of the beta-barrel '
    'including the wall atoms; they are therefore larger than the inner '
    'channel diameter (~24 \u00c5 in avGFP\u2078) reported in the original '
    'crystallographic papers, which refers to the solvent-accessible '
    'cavity through which the central alpha-helix passes. The beta-sheet '
    'walls add approximately 4\u20136 \u00c5 on each side, consistent with the '
    'mean minor axis of 30.4 \u00c5 reported here.'
)

add_subheading('Spectral and Photophysical Data.')
add_body(
    'Spectral data were obtained from FPbase\u00b3\u2074 using a hierarchical '
    'four-strategy matching protocol. (1) A manually curated dictionary '
    'mapped 176 PDB entries to known FPbase proteins. (2) The FPbase GraphQL '
    'API was queried by PDB identifier for entries with registered structures. '
    '(3) PDB titles were searched against a dictionary of over 150 fluorescent '
    'protein variant names (e.g., EGFP, mCherry, Venus, Dronpa) using '
    'case-insensitive keyword matching. (4) Remaining structures were assigned '
    'spectral properties by chromophore residue type where unambiguous. When '
    'multiple strategies returned matches, the highest-priority match was '
    'retained; each PDB identifier was mapped to at most one FPbase entry. '
    'Of 908 structures, 697 (76.8%) were matched to 71 unique FPbase protein '
    'names; 211 remained unmatched (predominantly structures with abbreviated '
    'or novel variant names not registered in FPbase). Quantum yield and '
    'extinction coefficient values from published '
    'sources\u00b2\u2070\u02c9\u00b3\u2078 were available for 316 '
    'structures. Color classes were assigned by emission '
    'wavelength: blue (<460 nm), cyan (460\u2013505 nm), green (505\u2013545 nm), '
    'yellow (545\u2013575 nm), orange (575\u2013610 nm), red (>610 nm). '
    'The complete mapping is provided in Table S1.'
)

add_subheading('B-Factor and Contact Analysis.')
add_body(
    'Mean B-factors were computed for chromophore and non-chromophore atoms '
    'separately. The B-factor ratio (chromophore/barrel) quantifies '
    'chromophore rigidity. Chromophore\u2013barrel contacts were counted as '
    'non-hydrogen barrel atoms within 4.0 \u00c5 of any chromophore atom. '
    'Chromophore atoms were defined as all non-hydrogen atoms belonging to '
    'the non-standard residue identified as the chromophore; this includes '
    'all atom types (imidazolinone ring, methine bridge, and pendant aromatic '
    'ring) for Tyr66-, Trp66-, and His66-derived chromophores alike.'
)

add_subheading('Statistical Analysis.')
add_body(
    'Analyses were performed in Python 3.14 using SciPy\u00b3\u2079 and '
    'scikit-learn.\u2074\u2070 Spearman rank correlations assessed monotonic '
    'relationships between continuous variables. Mann\u2013Whitney U tests '
    'compared two groups; Kruskal\u2013Wallis H tests compared multiple '
    'groups. All reported p-values were corrected for multiple testing using '
    'the Benjamini\u2013Hochberg false discovery rate (FDR) procedure; '
    'the family comprised 47 pre-specified hypothesis tests covering '
    'all pairwise Spearman correlations among the six geometry metrics, '
    'four photophysical variables, and resolution, plus group '
    'comparisons by color class and chromophore configuration. '
    'Raw and BH-corrected p-values are provided in Table S4. '
    'Partial correlations '
    'controlling for crystallographic resolution were computed by the '
    'residual method: both the dependent and independent variables were '
    'regressed on resolution using ordinary least squares, and the Spearman '
    'correlation of the residuals was reported. PCA of backbone C\u03b1 '
    'coordinates was used solely as a tool for determining the barrel axis '
    'direction (see Barrel Axis Determination); it was not used as an '
    'analytical technique for the photophysical data.'
)

add_subheading('AI-Assisted Computation.')
add_body(
    'All scripts for CIF parsing, coordinate transformation, convex hull '
    'computation, spectral matching, torsion angle calculation, and '
    'statistical analysis were generated by Claude (Anthropic, Opus 4) '
    'via Claude Code and reviewed by the authors before execution. The '
    'workflow followed a task-based structure similar to that described by '
    'Schwartz:\u00b3\u2070 each computational task was specified in plain '
    'language, the model produced code, and results were written to disk for '
    'independent verification. Failure modes included the model reporting '
    'steps as validated without performing checks, simplifying code based on '
    'patterns from unrelated examples, and halting after finding a single '
    'error rather than continuing to search. These were mitigated by '
    'requiring intermediate output files, cross-checking representative '
    'structures manually, and repeating verification prompts until no new '
    'errors were found. All scientific questions, methodology choices, and '
    'interpretations were made by the human authors.'
)

add_subheading('AlphaFold Structure Analysis.')
add_body(
    'AlphaFold v6 coordinate files were downloaded for fluorescent protein '
    'sequences via the EBI AlphaFold API '
    '(https://alphafold.ebi.ac.uk/api/prediction/{UniProt_ID}), which '
    'returns the current model version URL. Two analyses were performed. '
    'First, AlphaFold structures were obtained for all 346 UniProt '
    'accessions in the FPbase sequence list and matched to crystal '
    'structures in the main dataset via the UniProt REST API '
    '(https://rest.uniprot.org/uniprotkb/{ID}.json), yielding 51 '
    'proteins with both an AlphaFold model and a matched crystal '
    'structure. The highest-resolution crystal structure was selected '
    'per protein for paired comparison. Second, all '
    '347 unique FP sequences in the FPbase UniProt sequence list '
    '(accessed April 2026) were analysed; 346 unique UniProt identifiers '
    'were resolved (one entry lacked a UniProt accession), of which 327 '
    'yielded successful AlphaFold predictions and barrel analyses (17 had '
    'no available AlphaFold model, and 3 failed quality filters).\n'
    'Because AlphaFold structures contain no mature chromophore (the '
    'precursor residues Ser\u2013Tyr\u2013Gly remain uncyclized), the '
    'chromophore z-coordinate used to position the cross-sectional slice '
    'was set to z = 0 (the barrel centroid in the PCA frame). This '
    'approximation places the slice at the geometric midpoint of the '
    'barrel, which corresponds closely to the chromophore level in all '
    'known FP crystal structures. The axis computation matched the '
    'crystal-structure pipeline exactly: major and minor axes were derived '
    'as 4\u221aeigenvalue of the 2D covariance matrix of all non-hydrogen '
    'atoms in the \u00b12.0 \u00c5 slice. AlphaFold B-factor fields '
    'contain per-residue pLDDT confidence scores (0\u2013100) rather than '
    'crystallographic B-factors; these were recorded but not used in '
    'cross-structure comparisons with crystallographic B-factor ratios. '
    'Spectral data for the 347-sequence analysis were retrieved from the '
    'FPbase REST API and matched by UniProt accession.'
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════
# RESULTS AND DISCUSSION
# ═══════════════════════════════════════════════════════════
add_heading('Results and Discussion')

add_subheading('Dataset Composition.')
add_body(
    'The dataset comprises 908 fluorescent protein crystal structures: 843 '
    'with a validated mature chromophore and 65 without. The 843 chromophore-containing '
    'structures span 70 distinct residue types (CRO, CR2, NRQ, CRQ, GYS, PIA, and '
    'others). Spectral data from FPbase were matched to 697 structures. '
    'Quantum yield data were available for 316. The color class distribution '
    'was green (n = 474), red (n = 94), yellow (n = 69), cyan (n = 48), '
    'orange (n = 11), and blue (n = 1). Resolution ranged from 0.77 to '
    '3.60 \u00c5 (mean 1.86 \u00b1 0.47 \u00c5). Because the PDB contains '
    'multiple entries for the same or closely related proteins, the 697 '
    'spectrally matched structures correspond to only 71 unique FPbase '
    'protein names. The effective sample size for cross-protein comparisons '
    'is therefore substantially smaller than the total structure count, and '
    'robustness to pseudoreplication is assessed in Table S2.'
)

add_subheading('Barrel Geometry.')
add_body(
    'Cross-sectional geometry was computed for all 908 structures. Mean '
    'cross-sectional area was 818 \u00b1 194 \u00c5\u00b2 (range 330\u2013'
    '1975 \u00c5\u00b2). Mean eccentricity was 0.45 \u00b1 0.14, '
    'indicating moderately elliptical cross-sections. Circularity averaged '
    '0.92 \u00b1 0.03. Major and minor axes averaged 35.1 \u00b1 5.1 \u00c5 '
    'and 30.4 \u00b1 2.7 \u00c5, respectively. Barrel length averaged '
    '52.2 \u00b1 10.9 \u00c5. Area and eccentricity were not correlated '
    '(\u03c1 = 0.065, p = 0.055), indicating that barrel size and shape '
    'vary independently.'
)

add_subheading('Chromophore Maturation and Barrel Shape.')
add_body(
    'Structures with a validated mature chromophore (n = 843) were compared '
    'to those without (n = 65). Cross-sectional area did not differ '
    '(818 \u00b1 189 vs 773 \u00b1 229 \u00c5\u00b2, p = 0.14). '
    'Circularity, eccentricity, minor axis, and major axis likewise did not '
    'differ significantly (p = 0.20, 0.56, 0.93, and 0.62, respectively). '
    'Structures lacking a mature chromophore showed marginally longer barrels '
    '(54.8 \u00b1 12.3 vs 51.8 \u00b1 10.6 \u00c5, p = 0.018), '
    'consistent with these entries including photoconvertible or immature '
    'forms that may have crystallized with extended or less-ordered termini. '
    'Overall, the cross-sectional geometry of the barrel is largely '
    'insensitive to chromophore maturation state, suggesting that the '
    'barrel fold is established prior to chromophore cyclization. '
    'These results are broadly consistent with molecular dynamics simulations '
    'reporting only modest barrel contraction upon chromophore '
    'maturation,\u00b2\u2074\u02c9\u00b2\u2075 and with the role of '
    'conserved glycines in accommodating barrel assembly.\u00b2\u2076'
)

add_figure('fig1_chromophore_effect.png',
    'Figure 1. Effect of chromophore maturation on barrel geometry. '
    'Cross-sectional geometry does not differ significantly between '
    'chromophore-containing (n = 843) and chromophore-absent (n = 65) '
    'structures for area, circularity, eccentricity, or axis dimensions. '
    'Structures without a mature chromophore show marginally longer barrels '
    '(p = 0.018); all other metrics are not significantly different.')

add_subheading('Barrel Geometry by Emission Color Class.')
add_body(
    'Barrel geometry varied with emission color class. Eccentricity '
    '(Kruskal\u2013Wallis H = 67.9, p = 2.8 \u00d7 10\u207b\u00b9\u00b3), '
    'minor axis (H = 87.3, p = 2.5 \u00d7 10\u207b\u00b9\u2077), and '
    'circularity (H = 70.8, p = 6.8 \u00d7 10\u207b\u00b9\u2074) all '
    'differed across groups. Red fluorescent proteins had the most '
    'elliptical barrels and narrowest minor axes; green had the most '
    'circular. Emission wavelength correlated with eccentricity '
    '(\u03c1 = +0.210, p = 2.1 \u00d7 10\u207b\u2078), minor axis '
    '(\u03c1 = \u20130.290, p = 6.3 \u00d7 10\u207b\u00b9\u2075), and '
    'circularity (\u03c1 = \u20130.204, p = 5.3 \u00d7 10\u207b\u2078). '
    'Emission did not correlate with the major axis (\u03c1 = +0.039, '
    'p = 0.31) or barrel length (\u03c1 = \u20130.051, p = 0.18). '
    'All reported p-values were corrected for multiple testing using '
    'the Benjamini\u2013Hochberg procedure (q < 0.05); the key '
    'correlations remain significant after correction. '
    'The correlations of emission with eccentricity, minor axis, and '
    'circularity survive partial correlation analysis controlling for '
    'resolution. The major axis and barrel length are uninformative. '
    'These data are consistent with a red-shift associated with narrowing '
    'of the barrel in one dimension, rather than a general contraction. '
    'In DsRed-type red chromophores, an additional oxidation step extends '
    'the \u03c0-conjugated system by forming an acylimine group at the '
    'peptide bond N-terminal to the chromophore-forming tripeptide.\u00b9\u2076 '
    'The resulting larger chromophore occupies a different steric volume '
    'within the barrel pocket, which may contribute to the observed '
    'asymmetric narrowing; however, the causal direction of this relationship '
    'cannot be established from static crystal structures alone.'
)

add_figure('fig2_color_class.png',
    'Figure 2. Barrel geometry by emission color class. Mean (\u00b1 SEM) '
    'minor axis, eccentricity, circularity, and cross-sectional area. '
    'Red fluorescent proteins have the narrowest, most elliptical barrels. '
    'Emission wavelength scatter plots are shown in Figure S1.')

add_subheading('B-Factor Ratio and Quantum Yield.')
add_body(
    'The B-factor ratio (chromophore/barrel) was computed for 803 '
    'structures. The mean ratio was 0.80 \u00b1 0.20; in 87.8% of '
    'structures the chromophore was more rigid than the barrel (ratio < 1). '
    'The B-factor ratio was the strongest predictor of quantum yield: '
    '\u03c1 = \u20130.493, p = 7.9 \u00d7 10\u207b\u00b9\u2078, n = 268. '
    'This correlation survived partial correlation controlling for '
    'resolution (partial \u03c1 = \u20130.543). The ratio varied by color '
    'class: cyan 0.70 \u00b1 0.15, green 0.74 \u00b1 0.15, yellow '
    '0.79 \u00b1 0.16, orange 0.86 \u00b1 0.19, red 1.00 \u00b1 0.24. '
    'In red fluorescent proteins, the chromophore is on average no more '
    'rigid than the barrel.'
)

add_body(
    'The correlation between B-factor ratio and quantum yield is not simply '
    'a proxy for emission wavelength. Although emission correlates with the '
    'B-factor ratio (\u03c1 = +0.403), the B-factor ratio carries '
    'information about quantum yield beyond emission wavelength alone. The '
    'physical interpretation is direct: quantum yield depends on the '
    'competition between radiative and non-radiative decay. Non-radiative '
    'decay requires conformational motion, particularly rotation about the '
    'methine bridge. A chromophore that is rigid relative to its barrel has '
    'fewer accessible conformational states and therefore fewer decay '
    'pathways. This interpretation is consistent with the engineering of '
    'mTurquoise2,\u00b9\u2077 where a single packing mutation raised the '
    'quantum yield to 0.93, StayGold,\u2074\u00b9 which exhibits exceptional '
    'photostability, and YuzuFP,\u00b2\u2078 where altering residue 148 '
    'increased brightness 1.5-fold. The extended conjugation that produces '
    'red emission introduces torsional degrees of freedom that the barrel '
    'does not fully constrain, explaining the mean B-factor ratio of '
    '1.00 \u00b1 0.24 for red variants compared to 0.70 \u00b1 0.15 for '
    'cyan.'
)

add_figure('fig4_bfactor.png',
    'Figure 3. B-factor ratio analysis. (A) B-factor ratio vs quantum yield '
    '(\u03c1 = \u20130.493). (B) B-factor ratio vs emission wavelength. '
    '(C) B-factor ratio by color class; dashed line at 1.0 indicates equal '
    'rigidity.')

add_subheading('Chromophore\u2013Barrel Contacts.')
add_body(
    'The mean number of barrel atoms within 4.0 \u00c5 of the chromophore '
    'was 129 \u00b1 62 (n = 772). Contact count correlated with emission '
    'wavelength (\u03c1 = +0.182, p = 6.8 \u00d7 10\u207b\u2076, n = 607).'
)

add_subheading('Chromophore Dihedral Angles.')
add_body(
    'Megley \u03c4 and \u03c6 angles were computed for 690 structures. '
    'Cis configurations predominated (591, 85.7%), with 80 trans (11.6%) '
    'and 19 twisted (2.8%). The \u03c4 angle correlated with eccentricity '
    '(\u03c1 = \u20130.239, p = 2.0 \u00d7 10\u207b\u00b9\u2070). '
    'Absolute twist magnitudes correlated with eccentricity (|\u03c4|: '
    '\u03c1 = +0.196; |\u03c6|: \u03c1 = +0.205), quantum yield '
    '(|\u03c4|: \u03c1 = \u20130.239, p = 1.5 \u00d7 10\u207b\u2074; '
    '|\u03c6|: \u03c1 = \u20130.193, p = 0.0024), and B-factor ratio '
    '(|\u03c4|: \u03c1 = +0.117; |\u03c6|: \u03c1 = +0.143). The '
    'sum \u03c4 + \u03c6 was the strongest dihedral predictor '
    'of emission (\u03c1 = \u20130.304, p = 3.2 \u00d7 10\u207b\u00b9\u00b3) '
    'and quantum yield (\u03c1 = +0.270, p = 1.7 \u00d7 10\u207b\u2075). '
    'We note that \u03c4 + \u03c6 is used here as a descriptive algebraic '
    'sum of two static crystallographic dihedral angles, not in the '
    'mechanistic sense of the hula-twist photoisomerization pathway. '
    'This extends the findings of Megley et al.\u00b2\u2077 from a small '
    'set of structures to 690. Within this dataset, \u03c4 and \u03c6 are '
    'negatively correlated across all structures (\u03c1 = \u20130.411, '
    'p < 0.001; Figure 4A), forming two parallel bands in the \u03c4\u2013\u03c6 '
    'plane corresponding to cis and trans configurations. This anticorrelation '
    'primarily reflects the bimodal distribution of configurational states: '
    'cis structures (\u03c4 \u2248 0\u00b0) and trans structures '
    '(\u03c4 \u2248 \u00b1180\u00b0) each occupy distinct, non-overlapping '
    'regions of the \u03c4\u2013\u03c6 plane, such that the negative '
    'population-level correlation is a geometric consequence of the '
    'discrete classification rather than evidence for dynamic coupling '
    'between the two bonds in individual structures. Nevertheless, the '
    'relative positioning of \u03c4 and \u03c6 within each class may '
    'influence effective conjugation length. Twisted chromophores '
    'reside in more eccentric barrels with narrower minor axes. These '
    'observations are consistent with a model in which extended conjugation '
    'is associated with an asymmetric barrel that provides less symmetric '
    'constraint on torsional freedom, potentially allowing greater '
    'non-planar distortion and increased non-radiative decay. Trans-configured chromophores had higher '
    'eccentricity than cis (0.479 vs 0.443, p = 0.0023), lower circularity '
    '(0.913 vs 0.922, p = 0.0020), and narrower minor axes '
    '(29.9 vs 30.9 \u00c5, p = 1.0 \u00d7 10\u207b\u2075; Figure S2).'
)

add_figure('fig5_megley.png',
    'Figure 4. Chromophore dihedral analysis. (A) \u03c4 vs \u03c6 plot '
    'colored by emission class. (B) |\u03c4| vs quantum yield. (C) '
    'Dihedral sum (\u03c4 + \u03c6) vs emission wavelength.')

add_subheading('Resolution Confound and Barrel Size.')
add_body(
    'Resolution correlated with area (\u03c1 = \u20130.239, '
    'p = 7.3 \u00d7 10\u207b\u00b9\u00b3), eccentricity (\u03c1 = +0.148, '
    'p = 1.1 \u00d7 10\u207b\u2075), and minor axis (\u03c1 = \u20130.134, '
    'p = 7.2 \u00d7 10\u207b\u2075). Partial correlations controlling for '
    'resolution confirmed the key findings: emission vs eccentricity '
    '(partial \u03c1 = +0.206), emission vs minor axis '
    '(partial \u03c1 = \u20130.286), and QY vs B-factor ratio '
    '(partial \u03c1 = \u20130.543). Emission vs area weakened to '
    'partial \u03c1 = \u20130.129 (Figure S3). Cross-sectional area showed '
    'only weak correlation with emission (\u03c1 = \u20130.138), and barrel '
    'length was uninformative (\u03c1 = \u20130.051, p = 0.18). The '
    'photophysically relevant variation is in barrel shape, not size. This '
    'suggests that optimization of quantum yield should target the symmetry '
    'and tightness of chromophore packing in the minor-axis direction, not '
    'overall barrel volume.'
)

add_subheading('Stokes Shift and Barrel Geometry.')
add_body(
    'The Stokes shift\u2014the difference between excitation and emission '
    'maxima\u2014was available for 697 structures (mean 22.8 \u00b1 17.4 nm, '
    'range 9\u2013180 nm). Stokes shift varied by color class: cyan variants '
    'showed the largest mean Stokes shift (37.7 \u00b1 4.4 nm), followed by '
    'red (40.2 \u00b1 35.6 nm, with high variance), green (19.3 \u00b1 9.0 nm), '
    'orange (16.6 \u00b1 8.7 nm), and yellow (13.3 \u00b1 1.2 nm). Stokes '
    'shift correlated negatively with quantum yield '
    '(\u03c1 = \u20130.367, p = 1.7 \u00d7 10\u207b\u00b9\u00b9, n = 316), '
    'consistent with the interpretation that greater excited-state '
    'reorganization energy competes with radiative decay.'
)
add_body(
    'Stokes shift correlated significantly with barrel shape: proteins with '
    'larger Stokes shifts reside in more elliptical barrels '
    '(eccentricity: \u03c1 = +0.109, p = 0.004), with narrower minor axes '
    '(minor axis: \u03c1 = \u20130.111, p = 0.003), lower circularity '
    '(\u03c1 = \u20130.100, p = 0.008), and shorter barrel length '
    '(\u03c1 = \u20130.133, p = 4.4 \u00d7 10\u207b\u2074). Cross-sectional '
    'area was uninformative (\u03c1 = +0.011, p = 0.77). Partial correlations '
    'controlling for emission wavelength confirmed that these associations are '
    'not simply driven by the color-class dependence of Stokes shift: '
    'eccentricity (partial \u03c1 = +0.139, p = 2.4 \u00d7 10\u207b\u2074), '
    'minor axis (partial \u03c1 = \u20130.154, p = 4.5 \u00d7 10\u207b\u2075), '
    'circularity (partial \u03c1 = \u20130.128, p = 6.8 \u00d7 10\u207b\u2074), '
    'and barrel length (partial \u03c1 = \u20130.140, p = 2.0 \u00d7 10\u207b\u2074) '
    'each remained significant. The barrel length association is notable given '
    'that barrel length was uninformative for emission wavelength '
    '(\u03c1 = \u20130.051, p = 0.18), suggesting it captures variation '
    'relevant to excited-state reorganization independently of ground-state '
    'spectral tuning.'
)

add_subheading('Multivariate Analysis.')
add_body(
    'To disentangle the contributions of correlated predictors, multiple '
    'linear regression was performed with quantum yield as the response '
    'variable. A model including eccentricity, minor axis, B-factor ratio, '
    'and |\u03c4| yielded R\u00b2 = 0.43 (n = 244). The B-factor ratio was '
    'the only independently significant geometric predictor '
    '(\u03b2 = \u20130.096, p < 0.001); eccentricity, minor axis, and '
    '|\u03c4| were not significant after controlling for the other '
    'predictors. This indicates that the bivariate correlations between '
    'barrel shape and quantum yield are largely mediated through '
    'chromophore rigidity.'
)

add_body(
    'For emission wavelength, a model including minor axis, circularity, '
    'B-factor ratio, and resolution yielded R\u00b2 = 0.34 (n = 607). '
    'Minor axis (\u03b2 = \u20134.9, p = 4.6 \u00d7 10\u207b\u2074), '
    'circularity (\u03b2 = \u20138.5, p < 0.001), and B-factor ratio '
    '(\u03b2 = +14.8, p < 0.001) were each independently significant, '
    'indicating that barrel shape and chromophore rigidity contribute to '
    'emission wavelength through partially independent pathways.'
)

add_body(
    'A Random Forest model (500 trees, 5-fold cross-validation) predicting '
    'quantum yield achieved out-of-bag R\u00b2 = 0.57, with B-factor ratio '
    'as the dominant feature (importance = 0.563). The nonlinear model\u2019s '
    'improved performance suggests that interactions among predictors '
    'contribute to quantum yield determination beyond what linear models '
    'capture.'
)

add_subheading('Implications for Engineering.')
add_body(
    'These correlations do not establish causation, and the hypotheses '
    'below remain to be tested experimentally. Nevertheless, the observed '
    'trends suggest candidate strategies. If the B-factor ratio\u2013quantum '
    'yield correlation reflects an underlying causal relationship, then '
    'mutations that rigidify the chromophore relative to the barrel would '
    'be expected to increase quantum yield. Similarly, barrel eccentricity '
    'in a crystal structure might serve as a coarse diagnostic: high '
    'eccentricity could flag variants that are candidates for improvement, '
    'though other factors (chromophore chemistry, protonation state, '
    'excited-state dynamics) will also be important. The engineering '
    'successes of mTurquoise2,\u00b9\u2077 StayGold,\u2074\u00b9 and '
    'YuzuFP\u00b2\u2078\u2014each achieved through mutations altering '
    'chromophore\u2013barrel packing\u2014are consistent with these hypotheses '
    'but do not constitute a prospective test of them.'
)

add_subheading('AlphaFold Structural Predictions.')
add_body(
    'To assess whether AlphaFold-predicted structures recapitulate '
    'crystallographic barrel geometry, a paired analysis was performed '
    'for all 51 FP sequences in Maddie’s FPbase UniProt list for '
    'which both an AlphaFold model and a matched crystal structure exist '
    'in the dataset (UniProt–PDB cross-references from the UniProt '
    'REST API; highest-resolution crystal structure selected per protein; '
    'mean crystal resolution 1.63 ± 0.37 Å; mean AlphaFold pLDDT '
    '96.7 ± 1.3). Wilcoxon signed-rank tests showed that barrel '
    'shape metrics were indistinguishable between AlphaFold and crystal '
    'structures: eccentricity (Δ = −0.02, p = 0.21) and '
    'circularity (Δ = 0.00, p = 0.66) did not differ. However, '
    'AlphaFold predicted systematically narrower cross-sections: '
    'minor axis Δ = −0.53 Å (p = 0.002, 95% CI '
    '[−0.89, −0.15]) and major axis Δ = −0.89 Å '
    '(p = 0.002, 95% CI [−1.40, −0.41]). Cross-sectional area '
    'was 77.8 Å² smaller in AlphaFold structures '
    '(p < 0.001, 95% CI [−95.0, −61.0]). Barrel length was '
    '3.7 Å longer in AlphaFold predictions '
    '(p < 0.001, 95% CI [+1.5, +7.0]). '
    'The dimensional underestimate is consistent with the absence of a '
    'mature chromophore: the cyclized chromophore contributes steric '
    'bulk that slightly expands the barrel interior. The barrel length '
    'excess likely reflects AlphaFold modeling of flexible terminal '
    'regions that are disordered and excluded from electron density in '
    'crystal structures. Overall, AlphaFold faithfully captures barrel '
    'shape (eccentricity, circularity) but predicts slightly narrower '
    'cross-sections (~0.5–0.9 Å per axis) and longer barrels '
    '(+3.7 Å) compared to crystal structures.'
)

add_body(
    'Second, a population-level comparison was performed between all '
    '327 successfully analysed AlphaFold structures and the 843 '
    'chromophore-containing crystal structures. AlphaFold barrels '
    'showed smaller mean minor axis (28.6 ± 1.3 vs 30.5 ± 2.4 Å, '
    'Δ = −1.9 Å, p < 10⁻⁴⁹), higher eccentricity '
    '(0.51 ± 0.11 vs 0.45 ± 0.13, p < 10⁻¹⁴), lower '
    'circularity (0.91 ± 0.02 vs 0.92 ± 0.03, p < 10⁻²⁶), '
    'and substantially smaller area (698 ± 50 vs 818 ± 189 Å², '
    'p < 10⁻⁶⁴). Major axis did not differ significantly '
    '(33.8 ± 1.8 vs 35.0 ± 5.0 Å, p = 0.20). These '
    'population differences partly reflect the distinct protein '
    'compositions of the two datasets: the AlphaFold set includes '
    'many uncrystallized FPs that may be enriched in red-shifted, more '
    'eccentric variants, while the crystal set is dominated by '
    'well-studied green FPs. '
    'Notably, the correlation between emission maximum and minor axis '
    'length prominent in crystal structures (ρ = −0.29, p < 10⁻¹⁴) '
    'was absent in the AlphaFold ensemble (ρ = −0.001, p = 0.99). '
    'Instead, sequence length was the dominant structural correlate '
    '(ρ = −0.36 with eccentricity, p < 0.0001). This pattern '
    'indicates that the barrel distortions correlated with emission '
    'wavelength in crystal structures are driven by the mature '
    'chromophore—which is absent in AlphaFold models—rather than '
    'by primary sequence alone. AlphaFold therefore provides a useful '
    'geometric baseline for uncrystallized FPs, but reproducing the '
    'spectral correlations requires experimentally determined or '
    'computationally modeled chromophore geometry.'
)

add_subheading('Limitations.')
add_body(
    'Several caveats apply. The dataset is unevenly distributed: green '
    'proteins are overrepresented (n = 474), while blue (n = 1) and orange '
    '(n = 11) classes are sparse. The 908 PDB entries are not fully '
    'independent: many represent the same protein under different '
    'conditions, related mutants, or multiple chains. Of 697 entries matched '
    'to FPbase, only 71 unique protein names are represented. To assess '
    'pseudoreplication, the key correlations were recomputed after collapsing '
    'to one structure per unique protein (choosing the highest-resolution '
    'entry); the principal findings were preserved (Table S2).'
)

add_body(
    'The PDB is not a random sample of fluorescent proteins. Heavily '
    'studied variants such as GFP and mCherry are overrepresented, while '
    'many engineered or naturally occurring FPs have no deposited crystal '
    'structure. The correlations reported here therefore describe trends '
    'within the subset of FPs that have been crystallized, and their '
    'generalizability to the full diversity of fluorescent proteins remains '
    'to be established.'
)

add_body(
    'Crystal packing contacts are an uncontrolled variable in this '
    'dataset. Lattice contacts impose non-physiological forces on '
    'the barrel surface that can distort cross-sectional geometry, '
    'particularly eccentricity. Proteins that have been crystallized in '
    'multiple space groups (e.g., GFP variants, mCherry, DsRed) may '
    'show geometry variation attributable to different packing environments '
    'rather than intrinsic differences. Ideally, the analysis would be '
    'restricted to one crystal form per protein; the pseudoreplication '
    'collapse (Table S2) partially addresses this but does not fully '
    'control for packing effects.'
)

add_body(
    'The initial SCOP-based curation captured 877 structures; cross-referencing '
    'against all PDB entries annotated to FP-producing organisms (Aequorea, '
    'Discosoma, Entacmaea, Fungia, Galaxea, and related genera; total 1,295 '
    'entries) identified 31 additional standalone FP barrels deposited after '
    'the last SCOP update. These were subsequently analysed with the same '
    'pipeline and incorporated into the final dataset of 908 structures. The '
    'added structures include FusionRed (6U1A, 1.09 \u00c5), the mKate pH '
    'series (3BX9\u20133BXC), Kusabira-Cyan (2ZO6, 2ZO7), the darkmRuby pH '
    'series (7RHA\u20137RHC), photoswitchable Padron0.9 (3LS3, 3LSA), the '
    'Dreiklang photoswitchable GFP (3ST2\u20133ST4), AausFP2 (6S68), Azami '
    'Red1.0 (8I4K), PhoCl (7DMX, 7DNA), mNeptune684 (5YT1), mCrimson 0.9 '
    '(6WEM), pnRFP and B30Y variant (7LQO, 7LUG), LSSmOrange serial '
    'crystallography series (9LD5, 9LD8, 9LD9), DsRed-Monomer (8WGP), '
    'mCherry143azF (4ZIO), SAASoti variants (8PEI, 8ZBO), and sfGFP with '
    'unnatural amino acid (9S0T). An additional ~390 excluded entries were '
    'correctly filtered: 223 were solved by NMR or cryo-EM, and the remainder '
    'were FP fusions to non-FP domains (e.g., GCaMP2/3/6 calcium sensors in '
    'which circularly permuted GFP is fused to calmodulin, n = 23), '
    'multi-domain biosensors, or non-FP proteins from FP-producing organisms.'
)

add_body(
    'The use of crystallographic B-factors as a cross-structure rigidity '
    'metric requires caution. B-factors are influenced by resolution, '
    'temperature, refinement protocol, occupancy modeling, TLS treatment, '
    'and crystal packing, not solely by genuine atomic mobility. The '
    'B-factor ratio normalizes within each structure, mitigating some of '
    'these effects, and the correlation with quantum yield survives partial '
    'correlation controlling for resolution (partial \u03c1 = \u20130.543). '
    'Nevertheless, the B-factor ratio should be regarded as a proxy for '
    'relative chromophore mobility rather than a direct measure of dynamics.'
)

add_body(
    'The geometry pipeline depends on the choice of slice thickness '
    '(|z| \u2264 2.0 \u00c5), use of all non-hydrogen atoms, and convex hull '
    'construction. Sensitivity tests using slice thicknesses of 1.5 and '
    '2.5 \u00c5, and using backbone-heavy atoms only, confirmed that the '
    'rank order of structures by eccentricity, circularity, and minor axis '
    'length is preserved (Table S3). The convex hull does not capture '
    'internal cavities. Dihedral angles are static snapshots that do not '
    'capture solution-state dynamics. Quantum yield values were available '
    'for only 316 structures, from measurements under heterogeneous '
    'conditions. The full Spearman correlation matrix (Figure S4) and '
    'barrel geometry by chromophore type (Figure S5) are provided in '
    'Supporting Information.'
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════
add_heading('References')

refs = [
    '(1) Shimomura, O.; Johnson, F. H.; Saiga, Y. Extraction, Purification and Properties of Aequorin, a Bioluminescent Protein from the Luminous Hydromedusan, Aequorea. J. Cell. Comp. Physiol. 1962, 59, 223\u2013239.',
    '(2) Stepanenko, O. V.; Verkhusha, V. V.; Kuznetsova, I. M.; Uversky, V. N.; Turoverov, K. K. Fluorescent Proteins as Biomarkers and Biosensors: Throwing Color Lights on Molecular and Cellular Processes. Curr. Protein Pept. Sci. 2008, 9, 338\u2013369.',
    '(3) Chalfie, M.; Tu, Y.; Euskirchen, G.; Ward, W. W.; Prasher, D. C. Green Fluorescent Protein as a Marker for Gene Expression. Science 1994, 263, 802\u2013805.',
    '(4) The Nobel Foundation. The Nobel Prize in Chemistry 2008. https://www.nobelprize.org/prizes/chemistry/2008/ (accessed 2026-04-07).',
    '(5) Tsien, R. Y. The Green Fluorescent Protein. Annu. Rev. Biochem. 1998, 67, 509\u2013544.',
    '(6) Day, R. N.; Davidson, M. W. The Fluorescent Protein Palette: Tools for Cellular Imaging. Chem. Soc. Rev. 2009, 38, 2887\u20132921.',
    '(7) Rodriguez, E. A.; Campbell, R. E.; Lin, J. Y.; Lin, M. Z.; Miyawaki, A.; Palmer, A. E.; Shu, X.; Zhang, J.; Tsien, R. Y. The Growing and Glowing Toolbox of Fluorescent and Photoactive Proteins. Trends Biochem. Sci. 2017, 42, 111\u2013129.',
    '(8) Orm\u00f6, M.; Cubitt, A. B.; Kallio, K.; Gross, L. A.; Tsien, R. Y.; Remington, S. J. Crystal Structure of the Aequorea victoria Green Fluorescent Protein. Science 1996, 273, 1392\u20131395.',
    '(9) Yang, F.; Moss, L. G.; Phillips, G. N., Jr. The Molecular Structure of Green Fluorescent Protein. Nat. Biotechnol. 1996, 14, 1246\u20131251.',
    '(10) Heim, R.; Cubitt, A. B.; Tsien, R. Y. Improved Green Fluorescence. Nature 1995, 373, 663\u2013664.',
    '(11) Cubitt, A. B.; Heim, R.; Adams, S. R.; Boyd, A. E.; Gross, L. A.; Tsien, R. Y. Understanding, Improving and Using Green Fluorescent Proteins. Trends Biochem. Sci. 1995, 20, 448\u2013455.',
    '(12) Matz, M. V.; Fradkov, A. F.; Labas, Y. A.; Savitsky, A. P.; Zaraisky, A. G.; Markelov, M. L.; Lukyanov, S. A. Fluorescent Proteins from Nonbioluminescent Anthozoa Species. Nat. Biotechnol. 1999, 17, 969\u2013973.',
    '(13) Stepanenko, O. V.; Kuznetsova, I. M.; Verkhusha, V. V.; Turoverov, K. K. Beta-Barrel Scaffold of Fluorescent Proteins: Folding, Stability and Role in Chromophore Formation. Int. Rev. Cell Mol. Biol. 2013, 302, 221\u2013278.',
    '(14) Heim, R.; Tsien, R. Y. Engineering Green Fluorescent Protein for Improved Brightness, Longer Wavelengths and Fluorescence Resonance Energy Transfer. Curr. Biol. 1996, 6, 178\u2013182.',
    '(15) Wachter, R. M.; Elsliger, M. A.; Kallio, K.; Hanson, G. T.; Remington, S. J. Structural Basis of Spectral Shifts in the Yellow-Emission Variants of Green Fluorescent Protein. Structure 1998, 6, 1267\u20131277.',
    '(16) Gross, L. A.; Baird, G. S.; Hoffman, R. C.; Baldridge, K. K.; Tsien, R. Y. The Structure of the Chromophore within DsRed, a Red Fluorescent Protein from Coral. Proc. Natl. Acad. Sci. U.S.A. 2000, 97, 11990\u201311995.',
    '(17) Goedhart, J.; von Stetten, D.; Noirclerc-Savoye, M.; Lelimousin, M.; Joosen, L.; Hink, M. A.; van Weeren, L.; Gadella, T. W. J., Jr.; Royant, A. Structure-Guided Evolution of Cyan Fluorescent Proteins towards a Quantum Yield of 93%. Nat. Commun. 2012, 3, 751.',
    '(18) Ai, H. W.; Henderson, J. N.; Remington, S. J.; Campbell, R. E. Directed Evolution of a Monomeric, Bright and Photostable Version of Clavularia Cyan Fluorescent Protein. Biochem. J. 2006, 400, 531\u2013540.',
    '(19) Davidson, M. W.; Campbell, R. E. Engineered Fluorescent Proteins: Innovations and Applications. Nat. Methods 2009, 6, 713\u2013717.',
    '(20) Shaner, N. C.; Steinbach, P. A.; Tsien, R. Y. A Guide to Choosing Fluorescent Proteins. Nat. Methods 2005, 2, 905\u2013909.',
    '(21) Remington, S. J. Fluorescent Proteins: Maturation, Photochemistry and Photophysics. Curr. Opin. Struct. Biol. 2006, 16, 714\u2013721.',
    '(22) Craggs, T. D. Green Fluorescent Protein: Structure, Folding and Chromophore Maturation. Chem. Soc. Rev. 2009, 38, 2865\u20132875.',
    '(23) Zimmer, M. Green Fluorescent Protein (GFP): Applications, Structure, and Related Photophysical Behavior. Chem. Rev. 2002, 102, 759\u2013781.',
    '(24) Li, B.; Shahid, R.; Peshkepija, P.; Zimmer, M. Water Diffusion In and Out of the Beta-Barrel of GFP and the Fast Maturing Fluorescent Protein, TurboGFP. Chem. Phys. 2012, 392, 143\u2013148.',
    '(25) Zimmer, M. H.; Li, B.; Shahid, R. S.; Peshkepija, P.; Zimmer, M. Structural Consequences of Chromophore Formation and Exploration of Conserved Lid Residues amongst Naturally Occurring Fluorescent Proteins. Chem. Phys. 2014, 429, 5\u201311.',
    '(26) Nwafor, J.; Salguero, C.; Welcome, F.; Durmus, S.; Glasser, R. N.; Zimmer, M.; Schneider, T. L. Why Are Gly31, Gly33, and Gly35 Highly Conserved in All Fluorescent Proteins? Biochemistry 2021, 60, 3762\u20133770.',
    '(27) Megley, C. M.; Dickson, L. A.; Maddalo, S. L.; Chandler, G. J.; Zimmer, M. Photophysics and Dihedral Freedom of the Chromophore in Yellow, Blue, and Green Fluorescent Protein. J. Phys. Chem. B 2009, 113, 302\u2013308.',
    '(28) Ahmed, R. D.; Jamieson, W. D.; Vitsupakorn, D.; Zitti, A.; Pawson, K. A.; Castell, O. K.; Watson, P. D.; Jones, D. D. Molecular Dynamics Guided Identification of a Brighter Variant of Superfolder Green Fluorescent Protein with Increased Photobleaching Resistance. Commun. Chem. 2025, 8, 174.',
    '(29) Anthropic. Claude Code. https://docs.anthropic.com/en/docs/claude-code (accessed 2026-04-07).',
    '(30) Schwartz, M. D. Vibe Physics: The AI Grad Student. Anthropic 2026. https://www.anthropic.com/research/vibe-physics (accessed 2026-04-07).',
    '(31) Hartley, S. M.; Tiernan, K. A.; Ahmetaj, G.; Cretu, A.; Zhuang, Y.; Zimmer, M. AlphaFold2 and RoseTTAFold Predict Posttranslational Modifications. Chromophore Formation in GFP-like Proteins. PLoS One 2022, 17, e0267560.',
    '(32) Hayes, T.; et al. Simulating 500 Million Years of Evolution with a Language Model. Science 2025, 387, 850\u2013858.',
    '(33) Gujral, O.; Bafna, M.; Alm, E.; Berger, B. Sparse Autoencoders Uncover Biologically Interpretable Features in Protein Language Model Representations. Proc. Natl. Acad. Sci. U.S.A. 2025, 122, e2506316122.',
    '(34) Lambert, T. J. FPbase: A Community-Editable Fluorescent Protein Database. Nat. Methods 2019, 16, 277\u2013278.',
    '(35) Sievers, F.; et al. Fast, Scalable Generation of High-Quality Protein Multiple Sequence Alignments Using Clustal Omega. Mol. Syst. Biol. 2011, 7, 539.',
    '(36) Waterhouse, A. M.; Procter, J. B.; Martin, D. M. A.; Clamp, M.; Barton, G. J. Jalview Version 2\u2014A Multiple Sequence Alignment Editor and Analysis Workbench. Bioinformatics 2009, 25, 1189\u20131191.',
    '(37) Barber, C. B.; Dobkin, D. P.; Huhdanpaa, H. The Quickhull Algorithm for Convex Hulls. ACM Trans. Math. Softw. 1996, 22, 469\u2013483.',
    '(38) Cranfill, P. J.; et al. Quantitative Assessment of Fluorescent Proteins. Nat. Methods 2016, 13, 557\u2013562.',
    '(39) Virtanen, P.; et al. SciPy 1.0: Fundamental Algorithms for Scientific Computing in Python. Nat. Methods 2020, 17, 261\u2013272.',
    '(40) Pedregosa, F.; et al. Scikit-learn: Machine Learning in Python. J. Mach. Learn. Res. 2011, 12, 2825\u20132830.',
    '(41) Hirano, M.; Ando, R.; Shimozono, S.; Sugiyama, M.; Takeda, N.; Kurokawa, H.; Deguchi, R.; Endo, K.; Haga, K.; Takai-Todaka, R.; et al. A Highly Photostable and Bright Green Fluorescent Protein. Nat. Biotechnol. 2022, 40, 1132\u20131142.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    run = p.add_run(ref)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'

# ═══════════════════════════════════════════════════════════
# SUPPORTING INFORMATION
# ═══════════════════════════════════════════════════════════
doc.add_page_break()
add_heading('Supporting Information')

add_body(
    'The Supporting Information contains the following items:'
)

add_body(
    'Table S1. Complete dataset of 908 fluorescent protein crystal '
    'structures with PDB identifiers, chromophore types, cross-sectional '
    'geometry parameters, spectral data, B-factor ratios, dihedral angles, '
    'chromophore\u2013barrel contact counts, and photophysical properties. '
    'This table also serves as the source data for all figures in the '
    'manuscript (provided as a CSV file). Note: PDB entries 4DXQ and 4DXP '
    '(Kaede photoconvertible FP) are genuine FP barrels present in the PDB '
    'but absent from this dataset because their SCOP classification was not '
    'captured in our curation pipeline; they were not explicitly excluded on '
    'scientific grounds.'
)

add_body(
    'Table S2. Pseudoreplication robustness test. Key correlations and '
    'group comparisons recomputed after collapsing to one structure per '
    'unique FPbase protein name (n = 251, selecting the highest-resolution '
    'entry). Ten of eleven findings survive; the chromophore vs '
    'no-chromophore circularity comparison does not reach significance '
    'after collapsing (provided as a CSV file).'
)

add_body(
    'Table S3. Sensitivity analysis of the geometry pipeline. '
    'Cross-sectional parameters recomputed for 52 representative '
    'structures using slice thicknesses of 1.5 \u00c5 and 2.5 \u00c5 '
    '(vs default 2.0 \u00c5) and using backbone-heavy atoms only '
    '(vs all non-hydrogen atoms). Pearson and Spearman correlations '
    'with the default parameterization are reported (provided as a '
    'CSV file).'
)

add_body(
    'Note S1. AI Failure Case Study: Chromophore Detection. '
    'The chromophore detection step provides an instructive example of AI '
    'failure. The initial implementation used a hardcoded list of 31 residue '
    'names to identify mature chromophores. This list was incomplete in two '
    'directions. First, 62 additional chromophore residue codes present in the '
    'dataset were omitted, causing approximately 114 structures to be '
    'incorrectly classified as lacking a chromophore; among these were 1BFP '
    '(BFP-type chromophore IIC, His66-derived), 1EMF and 1EMK (CFP-type '
    'variants CSH and CCY, Trp66-derived), and others. Second, the residue '
    'code CR8 was present in the original list on the basis of its name '
    'resemblance to other chromophore codes (CR2, CR7, CRO), but the atom '
    'inventory check used standard Tyr66 atom names (CG2, CD1, CD2, CE1, '
    'CE2, CZ, OH) that are absent from CR8; CR8 instead uses a non-standard '
    'naming convention (C4\u2013C8, C11, C12, O13 for the phenol ring), so '
    'the aromatic ring was missed and 18 CR8-containing structures were '
    'incorrectly reclassified as non-chromophoric. Additional verified '
    'chromophores with non-standard atom naming (4F3, VYA) were caught by the '
    'same failure mode. Examination of the structures confirmed that CR8 '
    'contains both a complete imidazolinone ring and a para-hydroxyphenyl '
    'group bridged through the methine carbon (C8=CA2), consistent with a '
    'Tyr66-derived chromophore; the Arg residue at sequence position 66 in '
    'these entries corresponds to the conserved Arg96 in biological GFP '
    'numbering, which participates in chromophore formation but is not part '
    'of the chromophore itself. This class of error\u2014confidently wrong '
    'results arising from incomplete pattern matching\u2014is characteristic '
    'of LLM-generated code and underscores the necessity of human expert '
    'validation at each analytical step.'
)

add_body(
    'Code and Data Availability. All Python scripts used for CIF parsing, '
    'PCA-based barrel axis determination, cross-sectional slicing, '
    'convex hull analysis, spectral data matching, B-factor extraction, '
    'dihedral angle computation, statistical analysis, and figure '
    'generation were generated with the assistance of Claude (Anthropic) '
    'via Claude Code and are deposited in a public repository '
    '(https://github.com/[repository]; DOI: [to be assigned upon acceptance]). '
    'The complete dataset (Table S1) is provided as a CSV file in the '
    'Supporting Information.'
)

add_figure('fig3_emission_vs_geometry.png',
    'S1. Emission wavelength vs minor axis, eccentricity, and '
    'circularity. Each point represents one structure, colored by emission '
    'class.')

add_figure('fig9_cis_trans.png',
    'S2. Barrel geometry by chromophore configuration (cis vs trans).')

add_figure('fig6_resolution.png',
    'S3. Resolution control. (A) Resolution vs minor axis. '
    '(B) Emission vs minor axis colored by resolution. (C) High-resolution '
    'subset (<2.0 \u00c5) only.')

add_figure('fig7_heatmap.png',
    'S4. Spearman correlation matrix for all variables. Asterisks '
    'indicate * p < 0.05, ** p < 0.01, *** p < 0.001.')

add_figure('fig8_chromophore_types.png',
    'S5. Barrel geometry by chromophore residue type.')

# ── Save ──
output_path = '/Users/lukebegg/Downloads/GFP_JCIM_Manuscript.docx'
doc.save(output_path)
print(f"Saved: {output_path}")
